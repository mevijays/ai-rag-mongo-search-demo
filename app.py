import io
import os
from datetime import datetime

from flask import Flask, request, jsonify
from flask import Response
from dotenv import load_dotenv
from pymongo import MongoClient, ASCENDING
from werkzeug.utils import secure_filename

import numpy as np
from pypdf import PdfReader
from docx import Document as DocxDocument

from openai import OpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.vectorstores import VectorStore
from langchain_core.documents import Document

text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)

def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    denom = (np.linalg.norm(a) * np.linalg.norm(b))
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


class MongoNaiveVectorStore(VectorStore):
    def __init__(self, collection):
        self.collection = collection

    def add_texts(self, texts, metadatas=None, ids=None, **kwargs):
        if not texts:
            return []
        vecs = embed_texts(texts)
        now = datetime.utcnow()
        docs = []
        returned_ids = []
        for i, (t, v) in enumerate(zip(texts, vecs)):
            doc_id = ids[i] if ids and i < len(ids) else None
            meta = metadatas[i] if metadatas and i < len(metadatas) else {}
            rec = {
                "content": t,
                "embedding": v.tolist(),
                "created_at": now,
            }
            rec.update(meta or {})
            if doc_id is not None:
                rec["_id"] = doc_id
            docs.append(rec)
        if docs:
            result = self.collection.insert_many(docs)
            returned_ids = [str(x) for x in result.inserted_ids]
        return returned_ids

    def similarity_search_by_vector(self, embedding, k=4, fetch_k=200, **kwargs):
        docs = list(self.collection.find({}, {"content": 1, "title": 1, "source": 1, "chunk_index": 1, "embedding": 1}).limit(fetch_k))
        query_vec = np.array(embedding, dtype=np.float32)
        scored = []
        for d in docs:
            emb = np.array(d.get("embedding", []), dtype=np.float32)
            if emb.size == 0:
                continue
            scored.append((cosine_sim(query_vec, emb), d))
        scored.sort(key=lambda x: x[0], reverse=True)
        top = [d for _, d in scored[:k]]
        return [
            Document(page_content=d.get("content", ""), metadata={
                "title": d.get("title"),
                "source": d.get("source"),
                "chunk_index": d.get("chunk_index"),
            }) for d in top
        ]

    @classmethod
    def from_texts(cls, texts, metadatas=None, **kwargs):
        raise NotImplementedError("Use instance.add_texts; collection is required at init.")

    def similarity_search(self, query, k=4, **kwargs):
        q = openai_client.embeddings.create(model=EMBED_MODEL, input=[query]).data[0].embedding
        return self.similarity_search_by_vector(np.array(q, dtype=np.float32), k=k, **kwargs)

load_dotenv()

MONGO_HOST = os.getenv("MONGO_HOST", "localhost")
MONGO_PORT = int(os.getenv("MONGO_PORT", "27017"))
MONGO_USER = os.getenv("MONGO_INITDB_ROOT_USERNAME")
MONGO_PASS = os.getenv("MONGO_INITDB_ROOT_PASSWORD")
MONGO_DB_NAME = os.getenv("MONGO_DB_NAME", "ragdb")
VECTOR_COLLECTION = os.getenv("VECTOR_COLLECTION_NAME", "documents")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL")  # e.g., http://192.168.1.238:8000/v1
EMBED_MODEL = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")
CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")

FLASK_HOST = os.getenv("FLASK_HOST", "127.0.0.1")
FLASK_PORT = int(os.getenv("FLASK_PORT", "8000"))
FLASK_DEBUG = os.getenv("FLASK_DEBUG", "0") == "1"
RAG_TOP_K = int(os.getenv("RAG_TOP_K", "4"))
MAX_CONTEXT_CHARS = int(os.getenv("MAX_CONTEXT_CHARS", "10000"))

# Allow pointing to an OpenAI-compatible server via OPENAI_BASE_URL
# Many local servers ignore the API key but the SDK requires a value; provide a benign default if missing.
if OPENAI_BASE_URL:
    openai_client = OpenAI(api_key=OPENAI_API_KEY or "not-needed", base_url=OPENAI_BASE_URL)
else:
    openai_client = OpenAI(api_key=OPENAI_API_KEY)

mongo_uri = f"mongodb://{MONGO_USER}:{MONGO_PASS}@{MONGO_HOST}:{MONGO_PORT}/?authSource=admin"
mongo_client = MongoClient(mongo_uri)
db = mongo_client[MONGO_DB_NAME]
col = db[VECTOR_COLLECTION]

# Ensure indexes
col.create_index([("source", ASCENDING)])
col.create_index([("title", ASCENDING)])
col.create_index([("embedding", ASCENDING)], name="embedding_vector")

app = Flask(__name__)
vector_store = MongoNaiveVectorStore(col)

ALLOWED_EXTENSIONS = {"pdf", "docx"}

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def read_pdf(file_stream: io.BytesIO) -> str:
    reader = PdfReader(file_stream)
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(texts)

def read_docx(file_stream: io.BytesIO) -> str:
    document = DocxDocument(file_stream)
    return "\n".join(p.text for p in document.paragraphs)

def embed_texts(texts):
    if not texts:
        return []
    resp = openai_client.embeddings.create(model=EMBED_MODEL, input=texts)
    return [np.array(d.embedding, dtype=np.float32) for d in resp.data]

def retrieve_similar(query_vec: np.ndarray, k: int = 4):
    # Fetch a sample then rank by cosine in app; for demo simplicity
    docs = list(col.find({}, {"content": 1, "title": 1, "source": 1, "embedding": 1}).limit(200))
    scored = []
    for d in docs:
        emb = np.array(d.get("embedding", []), dtype=np.float32)
        if emb.size == 0:
            continue
        sim = cosine_sim(query_vec, emb)
        scored.append((sim, d))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [doc for _, doc in scored[:k]]

def llm_answer(prompt: str, context: str = "") -> str:
    # Build a compact prompt and avoid duplicating the context. Truncate to stay within model limits.
    if context:
        ctx = context if len(context) <= MAX_CONTEXT_CHARS else context[:MAX_CONTEXT_CHARS]
        user_content = (
            "Use the following context to answer. If the answer is not in the context, say 'not found in documents'.\n\n"
            f"Context:\n{ctx}\n\n"
            f"Question: {prompt}"
        )
    else:
        user_content = prompt

    messages = [{"role": "system", "content": "You are a helpful assistant."}, {"role": "user", "content": user_content}]
    chat = openai_client.chat.completions.create(model=CHAT_MODEL, messages=messages, temperature=0.2)
    return chat.choices[0].message.content.strip()

@app.route("/", methods=["GET"])
def index():
    html = """
    <!doctype html>
    <html>
    <head>
      <meta charset='utf-8'/>
      <title>RAG Demo</title>
      <style>
                :root { --border:#e5e7eb; --muted:#6b7280; --bg:#f9fafb; --primary:#2563eb; }
                body { font-family: -apple-system, Arial, sans-serif; margin: 24px; background: var(--bg); }
                h2 { margin-top: 0; }
                .card { background: white; border: 1px solid var(--border); padding: 16px; border-radius: 12px; box-shadow: 0 1px 2px rgba(0,0,0,.04); }
                .card.compact { padding: 12px; }
                .card h3 { margin: 0 0 8px 0; }
                .row { display: flex; flex-direction: column; gap: 24px; }
                .col { width: 100%; }
                textarea, input[type=text] { width: 100%; padding: 10px 12px; border:1px solid var(--border); border-radius:8px; }
                textarea { min-height: 96px; line-height: 1.4; resize: vertical; }
                button { padding: 10px 14px; border-radius: 8px; border:1px solid var(--border); background:white; cursor:pointer; }
                button.primary { background: var(--primary); color:white; border-color: var(--primary); }
                .meta { color: var(--muted); font-size: 12px; }
                .stack { display:flex; flex-direction:column; gap:12px; }
                .stack.tight { gap:8px; }
                .dropzone { border:2px dashed var(--border); padding:18px; border-radius:10px; text-align:center; color:var(--muted); background:#fff; }
                .dropzone.drag { border-color: var(--primary); color: var(--primary); background:#eef2ff; }
                .dropzone.small { border-width:1px; padding:10px; font-size:13px; }
                .answer { white-space: pre-wrap; padding:12px; border:1px solid var(--border); border-radius:10px; background:#fff; }
                .ctx { border-left:3px solid var(--border); padding-left:10px; margin:8px 0; }
                .badge { display:inline-block; padding:2px 8px; border-radius:999px; font-size:12px; border:1px solid var(--border); }
                .badge.warn { border-color:#f59e0b; color:#b45309; }
                .row-actions { display:flex; gap:8px; align-items:center; }
                .spinner { display:inline-block; width:16px; height:16px; border:2px solid var(--border); border-top-color: var(--primary); border-radius:50%; animation: spin 0.8s linear infinite; }
                @keyframes spin { to { transform: rotate(360deg); } }
      </style>
    </head>
    <body>
      <h2>Simple RAG with Mongo + OpenAI</h2>
      <div class='row'>
                <div class='col card compact' id='uploadCard'>
          <h3>Upload Document (PDF/DOCX)</h3>
                                        <div id='dz' class='dropzone small'>Drag & drop a PDF/DOCX here or click to choose</div>
                                        <form id='uploadForm' enctype='multipart/form-data' class='stack tight' style='margin-top:6px;'>
                                                <input id='fileInput' type='file' name='file' accept='.pdf,.docx' required style='display:none;' />
                                                <details>
                                                    <summary class='meta'>Options</summary>
                                                    <input type='text' name='title' placeholder='Optional title' />
                                                </details>
                        <div class='row-actions'>
                            <button id='uploadBtn' type='submit' class='primary'>Upload & Embed</button>
                            <span id='uploadSpin' style='display:none;' class='spinner'></span>
                        </div>
                    </form>
                    <div class='meta' id='uploadMeta'></div>
        </div>
        <div class='col card'>
          <h3>RAG Search</h3>
                    <div class='stack'>
                        <textarea id='q' placeholder='Ask a question... (Cmd/Ctrl+Enter to search)'></textarea>
                        <div class='row-actions'>
                            <button id='searchBtn' class='primary'>Search</button>
                            <span id='searchSpin' style='display:none;' class='spinner'></span>
                        </div>
                        <div id='fallbackBadge' class='badge warn' style='display:none;'>Used fallback LLM</div>
                        <div id='answer' class='answer'></div>
                        <button id='copyBtn' style='width:max-content;'>Copy Answer</button>
                        <div id='contexts'></div>
                    </div>
        </div>
      </div>

      <script>
                const dz = document.getElementById('dz');
                const fileInput = document.getElementById('fileInput');
                const uploadForm = document.getElementById('uploadForm');
                const uploadBtn = document.getElementById('uploadBtn');
                const uploadSpin = document.getElementById('uploadSpin');
                const uploadMeta = document.getElementById('uploadMeta');

                const searchBtn = document.getElementById('searchBtn');
                const searchSpin = document.getElementById('searchSpin');
                const qEl = document.getElementById('q');
                const answerEl = document.getElementById('answer');
                const contextsEl = document.getElementById('contexts');
                const fallbackBadge = document.getElementById('fallbackBadge');
                const copyBtn = document.getElementById('copyBtn');

                function setUploading(state) {
                    uploadBtn.disabled = state;
                    uploadSpin.style.display = state ? 'inline-block' : 'none';
                }

                function setSearching(state) {
                    searchBtn.disabled = state;
                    qEl.disabled = state;
                    searchSpin.style.display = state ? 'inline-block' : 'none';
                }

                dz.addEventListener('click', () => fileInput.click());
                dz.addEventListener('dragover', (e) => { e.preventDefault(); dz.classList.add('drag'); });
                dz.addEventListener('dragleave', () => dz.classList.remove('drag'));
                dz.addEventListener('drop', (e) => {
                    e.preventDefault(); dz.classList.remove('drag');
                    if (e.dataTransfer.files && e.dataTransfer.files[0]) {
                        fileInput.files = e.dataTransfer.files;
                    }
                });

                uploadForm.addEventListener('submit', async (e) => {
                    e.preventDefault();
                    uploadMeta.textContent = '';
                    if (!fileInput.files.length) { uploadMeta.textContent = 'Please choose a PDF or DOCX file.'; return; }
                    setUploading(true);
                    try {
                        const formData = new FormData(uploadForm);
                        const res = await fetch('/upload', { method: 'POST', body: formData });
                        const data = await res.json();
                        if (!res.ok) throw new Error(data.error || 'Upload failed');
                        uploadMeta.textContent = `Inserted ${data.inserted} chunks from "${data.title}" (total chunks: ${data.chunks}).`;
                        uploadForm.reset();
                        fileInput.value = '';
                    } catch (err) {
                        uploadMeta.textContent = `Error: ${err.message}`;
                    } finally {
                        setUploading(false);
                    }
                });

                async function doSearch() {
                    const q = qEl.value.trim();
                    if (!q) return;
                    setSearching(true);
                    answerEl.textContent = '';
                    contextsEl.innerHTML = '';
                    fallbackBadge.style.display = 'none';
                    try {
                        const res = await fetch('/search', { method: 'POST', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ q }) });
                        const data = await res.json();
                        if (!res.ok) throw new Error(data.error || 'Search failed');
                        answerEl.textContent = data.answer || '';
                        fallbackBadge.style.display = data.used_fallback_llm ? 'inline-block' : 'none';
                                    if (Array.isArray(data.context_docs)) {
                                        const pre = document.createElement('pre');
                                        pre.textContent = JSON.stringify(data.context_docs, null, 2);
                                        contextsEl.appendChild(pre);
                                    }
                    } catch (err) {
                        answerEl.textContent = `Error: ${err.message}`;
                    } finally {
                        setSearching(false);
                    }
                }

                searchBtn.addEventListener('click', doSearch);
            qEl.addEventListener('keydown', (e) => { if ((e.metaKey || e.ctrlKey) && e.key === 'Enter') doSearch(); });
                copyBtn.addEventListener('click', async () => {
                    try { await navigator.clipboard.writeText(answerEl.textContent || ''); copyBtn.textContent = 'Copied!'; setTimeout(()=> copyBtn.textContent='Copy Answer', 1000); } catch {}
                });
      </script>
    </body>
    </html>
    """
    return Response(html, mimetype="text/html")

@app.route("/upload", methods=["POST"])
def upload():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files['file']
    if f.filename == '':
        return jsonify({"error": "Empty filename"}), 400
    if not allowed_file(f.filename):
        return jsonify({"error": "Unsupported file type"}), 400

    filename = secure_filename(f.filename)
    ext = filename.rsplit('.', 1)[1].lower()
    raw = io.BytesIO(f.read())
    text = ""
    if ext == 'pdf':
        text = read_pdf(raw)
    elif ext == 'docx':
        text = read_docx(raw)

    if not text.strip():
        return jsonify({"error": "No text extracted from file"}), 400

    title = request.form.get('title') or filename
    chunks = text_splitter.split_text(text)
    metadatas = [{"title": title, "source": filename, "chunk_index": i} for i in range(len(chunks))]
    inserted_ids = vector_store.add_texts(chunks, metadatas=metadatas)
    return jsonify({"inserted": len(inserted_ids), "title": title, "chunks": len(chunks)})

@app.route("/search", methods=["POST"])
def search():
    data = request.get_json(force=True)
    query = data.get('q', '').strip()
    if not query:
        return jsonify({"error": "Query 'q' is required"}), 400

    qresp = openai_client.embeddings.create(model=EMBED_MODEL, input=[query])
    qvec = np.array(qresp.data[0].embedding, dtype=np.float32)
    top_docs = vector_store.similarity_search_by_vector(qvec, k=RAG_TOP_K)
    context = "\n---\n".join([d.page_content for d in top_docs])
    answer_from_docs = llm_answer(f"Answer the question based only on the context. If the answer isn't in the context, say 'not found in documents'.\nQuestion: {query}", context=context)

    if "not found in documents" in answer_from_docs.lower() or len(context.strip()) == 0:
        final_answer = llm_answer(query)
        used_fallback = True
    else:
        final_answer = answer_from_docs
        used_fallback = False

    return jsonify({
        "answer": final_answer,
        "used_fallback_llm": used_fallback,
        "context_docs": [{
            "title": d.metadata.get("title"),
            "source": d.metadata.get("source"),
            "chunk_index": d.metadata.get("chunk_index"),
            "snippet": (d.page_content or "")[:300]
        } for d in top_docs]
    })

if __name__ == "__main__":
    app.run(host=FLASK_HOST, port=FLASK_PORT, debug=FLASK_DEBUG)
